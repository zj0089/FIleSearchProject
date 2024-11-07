import os
import tkinter as tk
from tkinter import filedialog, messagebox, Scrollbar
import PyPDF2
import pdfplumber
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import docx
import pandas as pd
from openpyxl import Workbook
from typing import List, Dict
import re
from functools import lru_cache
import itertools
import unidecode

# Function to clear the cache of generate_name_variations
def clear_cache():
    generate_name_variations.cache_clear()

def search_names_in_text(text: str, variations: set) -> dict:
    results = {}
    text = text.lower()  # Normalize the text to lower case for case-insensitive matching
    for variation in variations:
        # Check if the variation consists of at least two parts (assumes parts are separated by spaces)
        if len(variation.split()) >= 2:
            pattern = re.compile(r'\b' + re.escape(variation) + r'\b', re.IGNORECASE)
            matches = pattern.findall(text)
            if matches:
                # Increment counts for each found variation
                found_variation = variation  # The variation as found in the text
                results[found_variation] = len(matches)
    return results

# Caching name variations for efficiency@lru_cache(maxsize=None)
@lru_cache(maxsize=None)
def generate_name_variations(name: str) -> set:
    # Normalize and split the name into parts, accounting for spaces and hyphens
    parts = [unidecode.unidecode(part.strip()) for part in re.split(r'[\s,;]+', name.replace('-', ' - ')) if part.strip()]
    variations = set()

    # Helper function to generate variations with initials, punctuations, and specific formats
    def add_variations(parts_list):
        full_variation = ' '.join(parts_list)
        variations.add(full_variation)
        if '-' in full_variation:
            variations.add(full_variation.replace(' - ', '-'))

        # Generate combinations with initials, avoiding all initials combo
        if len(parts_list) > 1:
            for index in range(len(parts_list)):
                # First part as initial followed by the rest
                if index == 0:
                    variations.add(parts_list[0][0] + '. ' + ' '.join(parts_list[1:]))
                # Last part as initial preceded by the rest
                elif index == len(parts_list) - 1:
                    variations.add(' '.join(parts_list[:-1]) + ' ' + parts_list[-1][0] + '.')
                # Middle parts as initials, with leading and trailing parts full
                else:
                    initial_variation = parts_list[:index] + [parts_list[index][0] + '.'] + parts_list[index+1:]
                    variations.add(' '.join(initial_variation))

                # Combinations of part with comma and initial (e.g., "Olmedo, L.")
                variations.add(parts_list[index] + ', ' + parts_list[0][0] + '.')

    # Specific combinations for names with exactly four parts
    if len(parts) == 4:
        add_variations([parts[0], parts[1]])  # 1st + 1st middle
        add_variations([parts[0], parts[2]])  # 1st + 2nd middle
        add_variations([parts[0], parts[3]])  # 1st + last
        add_variations([parts[0], parts[1], parts[3]])  # 1st + 1st middle + last
        add_variations([parts[0], parts[2], parts[3]])  # 1st + 2nd middle + last
    else:
        # Generate variations for names with less than 4 parts
        for num_parts in range(2, len(parts) + 1):
            for combo in itertools.combinations(parts, num_parts):
                add_variations(list(combo))

    return variations


# Extract text from PDF, optimizing for text, forms, and tables
def extract_text_from_pdf(file_path: str) -> Dict[int, str]:
    text_by_page = {}
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_by_page[page.page_number] = text
    return text_by_page

def extract_text_from_docx(file_path: str) -> Dict[int, str]:
    text_by_page = {}
    doc = docx.Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    text_by_page[1] = '\n'.join(text)
    return text_by_page

def extract_text_from_xlsx(file_path: str) -> Dict[int, str]:
    text_by_page = {}
    xls = pd.ExcelFile(file_path)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name)
        text_content = df.to_string(header=True, index=False)
        text_by_page[sheet_name] = text_content
    return text_by_page

# Function to search for names in files, updating to capture found variations
def search_names_in_files(folder_path: str, names_list: List[str]) -> Dict[str, List[Dict]]:
    results = {}
    names_variations = {name: generate_name_variations(name) for name in names_list}

    for root, dirs, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            if file.startswith('~$'):
                continue

            try:
                text_by_page = {}
                if file.endswith('.pdf'):
                    text_by_page = extract_text_from_pdf(file_path)
                elif file.endswith('.docx'):
                    text_by_page = extract_text_from_docx(file_path)
                elif file.endswith('.xlsx'):
                    text_by_page = extract_text_from_xlsx(file_path)

                for page_num, page_text in text_by_page.items():
                    for name, variations in names_variations.items():
                        found_results = search_names_in_text(page_text, variations)
                        for found_variation, occurrences in found_results.items():
                            results.setdefault(name, []).append({
                                'folder_name': os.path.basename(root),
                                'folder_path': root,
                                'file': file,
                                'type': file.split('.')[-1].upper(),
                                'page': page_num,
                                'occurrences': occurrences,
                                'variation': found_variation
                            })

            except Exception as e:
                print(f"Error processing {file_path}: {str(e)}")

    return results


# Save results to Excel using openpyxl
def save_results_to_excel(results: Dict[str, List[Dict]], output_file: str):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Search Results"
    headers = ['Name', 'Name Variation', 'Folder Name', 'Folder Path', 'File', 'Type', 'Sheet', 'Page', 'Occurrences']
    sheet.append(headers)
    for name, entries in results.items():
        for entry in entries:
            row = [
                name,
                entry['variation'],
                entry['folder_name'],
                entry['folder_path'],
                entry['file'],
                entry['type'],
                entry.get('sheet', ''),
                entry.get('page', ''),
                entry['occurrences']
            ]
            sheet.append(row)
    workbook.save(output_file)

# GUI application
def run_gui_app():
    def browse_folder():
        folder = filedialog.askdirectory()
        folder_path_entry.delete(0, tk.END)
        folder_path_entry.insert(0, folder)

    def load_names():
        file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt"), ("Excel files", "*.xlsx")])
        if file_path:
            if file_path.endswith('.txt'):
                with open(file_path, 'r') as file:
                    names = file.read().splitlines()
                    names_text.delete('1.0', tk.END)
                    names_text.insert(tk.END, "\n".join(names))
            elif file_path.endswith('.xlsx'):
                df = pd.read_excel(file_path)
                names = df.iloc[:, 0].tolist()  # Assuming names are in the first column
                names_text.delete('1.0', tk.END)
                names_text.insert(tk.END, "\n".join(names))

    def search_files():
        folder_path = folder_path_entry.get()
        names_list = names_text.get("1.0", tk.END).strip()
        
        # Split names by both newlines and commas, then strip whitespace
        names_list = [name.strip() for name in names_list.replace(',', '\n').splitlines() if name.strip()]

        if not folder_path or not names_list:
            messagebox.showwarning("Input Error", "Please provide a folder path and a list of names.")
            return

        results = search_names_in_files(folder_path, names_list)
        if results:
            result_text.delete("1.0", tk.END)
            for name, entries in results.items():
                result_text.insert(tk.END, f"{name} found in:\n")
                for entry in entries:
                    result_text.insert(
                        tk.END,
                        f"  - Folder: {entry['folder_name']} (Path: {entry['folder_path']}), File: {entry['file']} "
                        f"(Type: {entry['type']}, Sheet: {entry.get('sheet', '')}, Page: {entry.get('page', '')}, "
                        f"Row: {entry.get('row', '')}, Column: {entry.get('column', '')}): {entry['occurrences']} occurrence(s)\n"
                    )
                result_text.insert(tk.END, "\n")
            save_results_to_excel(results, 'search_results.xlsx')  # Save results to Excel
        else:
            result_text.delete("1.0", tk.END)
            result_text.insert(tk.END, "No matches found.")

    # Set up the main window
    window = tk.Tk()
    window.title("File Search Application")
    window.geometry("800x600")

    # Folder selection
    tk.Label(window, text="Select Folder Path:").pack(anchor="w", padx=10)
    folder_frame = tk.Frame(window)
    folder_frame.pack(anchor="w", padx=10, pady=5)
    folder_path_entry = tk.Entry(folder_frame, width=50)
    folder_path_entry.pack(side="left", fill="x", expand=True)
    tk.Button(folder_frame, text="Browse", command=browse_folder).pack(anchor="w", padx=5, pady=5, side="left")

    # Names input with vertical scroll bar
    tk.Label(window, text="Enter Names (one per line) or Load from File:").pack(anchor="w", padx=10)
    names_frame = tk.Frame(window)
    names_frame.pack(anchor="w", padx=10, pady=5)
    names_text = tk.Text(names_frame, height=10, width=60, wrap="word")
    names_text.pack(side="left", fill="both", expand=True)
    names_v_scroll = Scrollbar(names_frame, orient="vertical", command=names_text.yview)
    names_v_scroll.pack(side="right", fill="y")
    names_text.config(yscrollcommand=names_v_scroll.set)

    # Load names button
    tk.Button(window, text="Load Names from File", command=load_names).pack(anchor="w", padx=10, pady=5)

    # Search button
    tk.Button(window, text="Search Files", command=search_files, bg="green", fg="white").pack(anchor="w", padx=10, pady=10)

    # Results display with vertical scroll bar
    tk.Label(window, text="Results:").pack(anchor="w", padx=10)
    results_frame = tk.Frame(window)
    results_frame.pack(anchor="w", padx=10, pady=5, fill="both", expand=True)
    result_text = tk.Text(results_frame, wrap="word")
    result_text.pack(side="left", fill="both", expand=True)
    results_v_scroll = Scrollbar(results_frame, orient="vertical", command=result_text.yview)
    results_v_scroll.pack(side="right", fill="y")
    result_text.config(yscrollcommand=results_v_scroll.set)

    window.mainloop()

# Run the application
run_gui_app()
